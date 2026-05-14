import uuid
from io import BytesIO
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor
from fastapi import APIRouter, Depends, Form, HTTPException, Query, UploadFile, File, status
from fastapi.responses import StreamingResponse
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.auth import get_current_user
from app.database import get_db
from app.models.contract import Contract, ContractRevision
from app.models.evaluation import Evaluation, ProposalScore
from app.models.procurement import Procurement
from app.models.proposal import SupplierProposal
from app.models.rfp import RFP, RFPRevision
from app.models.user import User
from app.schemas.procurement import ProcurementCreate, ProcurementList, ProcurementOut, ProcurementUpdate

router = APIRouter(prefix="/procurements", tags=["procurements"])


@router.post("", response_model=ProcurementOut, status_code=status.HTTP_201_CREATED)
async def create_procurement(
    payload: ProcurementCreate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    procurement = Procurement(**payload.model_dump(), created_by=current_user.id)
    db.add(procurement)
    await db.commit()
    await db.refresh(procurement)
    return ProcurementOut.model_validate(procurement)


@router.get("", response_model=ProcurementList)
async def list_procurements(
    skip: int = Query(0, ge=0),
    limit: int = Query(20, ge=1, le=100),
    status: Optional[str] = Query(None),
    stage: Optional[str] = Query(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    query = select(Procurement).where(Procurement.created_by == current_user.id)
    if status:
        query = query.where(Procurement.status == status)
    if stage:
        query = query.where(Procurement.stage == stage)

    count_result = await db.execute(select(func.count()).select_from(query.subquery()))
    total = count_result.scalar()
    result = await db.execute(query.order_by(Procurement.updated_at.desc()).offset(skip).limit(limit))
    items = result.scalars().all()
    return ProcurementList(items=[ProcurementOut.model_validate(p) for p in items], total=total)


@router.get("/{procurement_id}", response_model=ProcurementOut)
async def get_procurement(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Procurement).where(
            Procurement.id == procurement_id, Procurement.created_by == current_user.id
        )
    )
    procurement = result.scalar_one_or_none()
    if not procurement:
        raise HTTPException(status_code=404, detail="Procurement not found")
    return ProcurementOut.model_validate(procurement)


@router.patch("/{procurement_id}", response_model=ProcurementOut)
async def update_procurement(
    procurement_id: uuid.UUID,
    payload: ProcurementUpdate,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    result = await db.execute(
        select(Procurement).where(
            Procurement.id == procurement_id, Procurement.created_by == current_user.id
        )
    )
    procurement = result.scalar_one_or_none()
    if not procurement:
        raise HTTPException(status_code=404, detail="Procurement not found")
    for field, value in payload.model_dump(exclude_none=True).items():
        setattr(procurement, field, value)
    await db.commit()
    await db.refresh(procurement)
    return ProcurementOut.model_validate(procurement)


@router.get("/{procurement_id}/proposals")
async def list_proposals(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    result = await db.execute(
        select(SupplierProposal)
        .where(SupplierProposal.procurement_id == procurement_id)
        .order_by(SupplierProposal.upload_timestamp.asc())
    )
    proposals = result.scalars().all()
    return [_proposal_dict(p) for p in proposals]


@router.post("/{procurement_id}/proposals", status_code=status.HTTP_201_CREATED)
async def upload_proposal(
    procurement_id: uuid.UUID,
    supplier_name: str = Form(...),
    text_content: Optional[str] = Form(None),
    file: Optional[UploadFile] = File(None),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)

    if not text_content and not file:
        raise HTTPException(status_code=400, detail="Provide either text_content or a file")

    upload_dir = Path(settings.UPLOAD_DIR) / str(procurement_id)
    upload_dir.mkdir(parents=True, exist_ok=True)

    file_path = ""
    mime_type = "text/plain"
    original_filename = f"{supplier_name}_proposal.txt"
    file_size = None

    if file and file.filename:
        original_filename = file.filename
        mime_type = file.content_type or "application/octet-stream"
        dest = upload_dir / f"{uuid.uuid4()}_{file.filename}"
        content = await file.read()
        file_size = len(content)
        dest.write_bytes(content)
        file_path = str(dest)

        raw_text = None
        if "text" in mime_type:
            raw_text = content.decode("utf-8", errors="replace")
    else:
        # Text submission — write to a .txt file so the workflow can read it
        dest = upload_dir / f"{uuid.uuid4()}_{supplier_name.replace(' ', '_')}.txt"
        raw_text = text_content
        dest.write_text(text_content, encoding="utf-8")
        file_path = str(dest)
        file_size = len(text_content.encode())

    proposal = SupplierProposal(
        procurement_id=procurement_id,
        supplier_name=supplier_name,
        original_filename=original_filename,
        file_path=file_path,
        file_size_bytes=file_size,
        mime_type=mime_type,
        raw_text=raw_text if not file or "text" in mime_type else None,
        extraction_status="SUBMITTED",
        status="SUBMITTED",
    )
    db.add(proposal)
    await db.commit()
    await db.refresh(proposal)
    return _proposal_dict(proposal)


@router.delete("/{procurement_id}/proposals/{proposal_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_proposal(
    procurement_id: uuid.UUID,
    proposal_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    result = await db.execute(
        select(SupplierProposal).where(
            SupplierProposal.id == proposal_id,
            SupplierProposal.procurement_id == procurement_id,
        )
    )
    proposal = result.scalar_one_or_none()
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")
    await db.delete(proposal)
    await db.commit()


async def _get_procurement_or_404(
    procurement_id: uuid.UUID, user_id: uuid.UUID, db: AsyncSession
) -> Procurement:
    result = await db.execute(
        select(Procurement).where(
            Procurement.id == procurement_id, Procurement.created_by == user_id
        )
    )
    p = result.scalar_one_or_none()
    if not p:
        raise HTTPException(status_code=404, detail="Procurement not found")
    return p


@router.get("/{procurement_id}/rfp")
async def get_rfp(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    result = await db.execute(
        select(RFP)
        .where(RFP.procurement_id == procurement_id)
        .order_by(RFP.version.desc())
        .limit(1)
    )
    rfp = result.scalar_one_or_none()
    if not rfp:
        raise HTTPException(status_code=404, detail="No RFP found")
    return {
        "id": str(rfp.id),
        "version": rfp.version,
        "status": rfp.status,
        "content": rfp.content,
        "model_id": rfp.model_id,
        "created_at": rfp.created_at.isoformat() if rfp.created_at else None,
    }


@router.get("/{procurement_id}/rfp/download")
async def download_rfp(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    procurement = await _get_procurement_or_404(procurement_id, current_user.id, db)
    result = await db.execute(
        select(RFP)
        .where(RFP.procurement_id == procurement_id)
        .order_by(RFP.version.desc())
        .limit(1)
    )
    rfp = result.scalar_one_or_none()
    if not rfp:
        raise HTTPException(status_code=404, detail="No RFP found")

    c = rfp.content or {}
    doc = Document()

    title = doc.add_heading("Request for Proposal", level=0)
    title.runs[0].font.color.rgb = RGBColor(0, 29, 61)

    sub = doc.add_paragraph(f"{procurement.title}   |   v{rfp.version}   |   Status: {rfp.status}")
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(80, 80, 80)
    doc.add_paragraph("")

    def add_section(heading: str, text: str):
        doc.add_heading(heading, level=1)
        for para in (text or "").split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    add_section("Executive Summary", c.get("executive_summary", ""))
    add_section("Scope of Work", c.get("scope_of_work", ""))

    deliverables = c.get("deliverables") or []
    if deliverables:
        doc.add_heading("Deliverables", level=1)
        for item in deliverables:
            doc.add_paragraph(str(item), style="List Bullet")

    add_section("Submission Requirements", c.get("submission_requirements", ""))

    criteria = c.get("evaluation_criteria") or []
    if criteria:
        doc.add_heading("Evaluation Criteria", level=1)
        for ec in criteria:
            criterion = ec.get("criterion", "") if isinstance(ec, dict) else str(ec)
            weight = ec.get("weight", 0) if isinstance(ec, dict) else 0
            doc.add_paragraph(f"• {criterion} — {round(float(weight) * 100)}%")

    timelines = c.get("timelines") or []
    if timelines:
        doc.add_heading("Timeline & Milestones", level=1)
        for ms in timelines:
            desc = ms.get("description", "") if isinstance(ms, dict) else str(ms)
            due = ms.get("due_date", "") if isinstance(ms, dict) else ""
            doc.add_paragraph(f"• {desc}" + (f"  ({due})" if due else ""))

    add_section("Legal & Compliance", c.get("legal_compliance_notes", ""))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_title = procurement.title.replace(" ", "_").replace("/", "_")[:40]
    filename = f"RFP_{safe_title}_v{rfp.version}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{procurement_id}/rfp/revisions")
async def get_rfp_revisions(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    rfp_result = await db.execute(
        select(RFP).where(RFP.procurement_id == procurement_id).order_by(RFP.version.desc()).limit(1)
    )
    rfp = rfp_result.scalar_one_or_none()
    if not rfp:
        return []
    rev_result = await db.execute(
        select(RFPRevision).where(RFPRevision.rfp_id == rfp.id).order_by(RFPRevision.version.asc())
    )
    return [
        {
            "id": str(r.id),
            "version": r.version,
            "revision_request": r.revision_request,
            "content": r.content,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in rev_result.scalars().all()
    ]


@router.get("/{procurement_id}/contract/revisions")
async def get_contract_revisions(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    contract_result = await db.execute(
        select(Contract).where(Contract.procurement_id == procurement_id).order_by(Contract.version.desc()).limit(1)
    )
    contract = contract_result.scalar_one_or_none()
    if not contract:
        return []
    rev_result = await db.execute(
        select(ContractRevision)
        .where(ContractRevision.contract_id == contract.id)
        .order_by(ContractRevision.version.asc())
    )
    return [
        {
            "id": str(r.id),
            "version": r.version,
            "revision_request": r.revision_request,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        for r in rev_result.scalars().all()
    ]


@router.get("/{procurement_id}/evaluation")
async def get_evaluation(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    eval_result = await db.execute(
        select(Evaluation)
        .where(Evaluation.procurement_id == procurement_id)
        .order_by(Evaluation.evaluated_at.desc().nulls_first())
        .limit(1)
    )
    evaluation = eval_result.scalar_one_or_none()
    if not evaluation:
        raise HTTPException(status_code=404, detail="No evaluation found")

    scores_result = await db.execute(
        select(ProposalScore, SupplierProposal.supplier_name)
        .join(SupplierProposal, ProposalScore.proposal_id == SupplierProposal.id)
        .where(ProposalScore.evaluation_id == evaluation.id)
        .order_by(ProposalScore.rank.asc())
    )
    rows = scores_result.all()

    return {
        "id": str(evaluation.id),
        "status": evaluation.status,
        "evaluation_weights": evaluation.evaluation_weights,
        "recommendation_proposal_id": str(evaluation.recommendation_proposal_id) if evaluation.recommendation_proposal_id else None,
        "recommendation_rationale": evaluation.recommendation_rationale,
        "scores": [
            {
                "proposal_id": str(row.ProposalScore.proposal_id),
                "supplier_name": row.supplier_name,
                "weighted_total": float(row.ProposalScore.weighted_total),
                "rank": row.ProposalScore.rank,
                "ai_assessment": row.ProposalScore.ai_assessment,
            }
            for row in rows
        ],
    }


@router.get("/{procurement_id}/contract")
async def get_contract(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    result = await db.execute(
        select(Contract)
        .where(Contract.procurement_id == procurement_id)
        .order_by(Contract.version.desc())
        .limit(1)
    )
    contract = result.scalar_one_or_none()
    if not contract:
        raise HTTPException(status_code=404, detail="No contract found")
    return {
        "id": str(contract.id),
        "supplier_name": contract.supplier_name,
        "version": contract.version,
        "draft_content": contract.draft_content,
        "status": contract.status,
        "model_id": contract.model_id,
        "created_at": contract.generation_timestamp.isoformat() if contract.generation_timestamp else None,
    }


_CONTRACT_SECTION_LABELS = {
    "scope": "Scope of Work",
    "payment_terms": "Payment Terms",
    "milestones": "Milestones",
    "legal_clauses": "Legal Clauses",
    "termination_clauses": "Termination",
}


@router.get("/{procurement_id}/contract/download")
async def download_contract(
    procurement_id: uuid.UUID,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    await _get_procurement_or_404(procurement_id, current_user.id, db)
    result = await db.execute(
        select(Contract)
        .where(Contract.procurement_id == procurement_id)
        .order_by(Contract.version.desc())
        .limit(1)
    )
    contract = result.scalar_one_or_none()
    if not contract:
        raise HTTPException(status_code=404, detail="No contract found")

    doc = Document()

    title = doc.add_heading("Procurement Contract", level=0)
    title.runs[0].font.color.rgb = RGBColor(0, 29, 61)

    sub = doc.add_paragraph(f"Supplier: {contract.supplier_name}   |   Version {contract.version}")
    sub.runs[0].font.size = Pt(11)
    sub.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")

    ordered_keys = [k for k in _CONTRACT_SECTION_LABELS if k in contract.draft_content]
    extras = [k for k in contract.draft_content if k not in _CONTRACT_SECTION_LABELS and k != "supplier_name"]
    for key in ordered_keys + extras:
        content = contract.draft_content[key]
        label = _CONTRACT_SECTION_LABELS.get(key, key.replace("_", " ").title())
        doc.add_heading(label, level=1)
        text = content if isinstance(content, str) else str(content)
        for para in text.split("\n\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    safe_name = contract.supplier_name.replace(" ", "_").replace("/", "_")
    filename = f"contract_{safe_name}_v{contract.version}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _proposal_dict(p: SupplierProposal) -> dict:
    return {
        "id": str(p.id),
        "procurement_id": str(p.procurement_id),
        "supplier_name": p.supplier_name,
        "original_filename": p.original_filename,
        "file_size_bytes": p.file_size_bytes,
        "mime_type": p.mime_type,
        "extraction_status": p.extraction_status,
        "status": p.status,
        "upload_timestamp": p.upload_timestamp.isoformat() if p.upload_timestamp else None,
        "extracted_data": p.extracted_data,
    }
